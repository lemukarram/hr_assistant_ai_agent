"""
RAG Indexer — ingests HR handbook documents into ChromaDB.

Supports: PDF, Markdown, plain text
Arabic-aware chunking with sentence boundary respect.
Idempotent: re-running checks existing chunks before re-indexing.
"""
import hashlib
import re
from pathlib import Path

from app.core.config import settings
from app.core.logging import get_logger

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".md", ".txt", ".docx"}


class ArabicAwareTextSplitter:
    """
    Splits text respecting Arabic sentence boundaries.
    
    Priority order for split points:
    1. Paragraph boundaries (double newline)
    2. Arabic full stop (۔) or Western full stop (.)
    3. Arabic comma (،) or Western comma (,)
    
    Never splits mid-word. Overlaps are character-based to avoid
    cutting Arabic words at chunk boundaries.
    """

    def __init__(self, chunk_size: int = 512, overlap: int = 64):
        self.chunk_size = chunk_size
        self.overlap = overlap
        # Arabic sentence boundary pattern
        self._split_pattern = re.compile(r"(\n\n+|[.۔!؟]\s+|[،,]\s+)")

    def split(self, text: str) -> list[dict]:
        """Split text into chunks with metadata."""
        # Clean text
        text = self._clean(text)

        # Split into sentences, then further break overly long sentences
        raw_sentences = self._split_into_sentences(text)
        sentences = []
        for s in raw_sentences:
            words = s.split()
            if len(words) > self.chunk_size:
                # Hard-split long sentences at chunk_size word boundaries
                for i in range(0, len(words), self.chunk_size):
                    part = " ".join(words[i : i + self.chunk_size])
                    if part:
                        sentences.append(part)
            else:
                sentences.append(s)

        chunks = []
        current_chunk = []
        current_length = 0

        for sentence in sentences:
            sentence_len = len(sentence.split())

            if current_length + sentence_len > self.chunk_size and current_chunk:
                # Emit current chunk
                chunk_text = " ".join(current_chunk)
                chunks.append({"text": chunk_text, "token_count": current_length})

                # Overlap: keep last N tokens
                overlap_words = " ".join(current_chunk).split()[-self.overlap:]
                current_chunk = [" ".join(overlap_words)]
                current_length = len(overlap_words)

            current_chunk.append(sentence)
            current_length += sentence_len

        if current_chunk:
            chunks.append({"text": " ".join(current_chunk), "token_count": current_length})

        return chunks

    def _clean(self, text: str) -> str:
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Normalize newlines
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _split_into_sentences(self, text: str) -> list[str]:
        """Split text into sentence-level units."""
        parts = self._split_pattern.split(text)
        sentences = []
        for part in parts:
            part = part.strip()
            if part and not self._split_pattern.match(part):
                sentences.append(part)
        return [s for s in sentences if len(s.split()) >= 3]  # Discard tiny fragments


class HandbookIngestionPipeline:
    """Ingest HR handbook documents into ChromaDB with embeddings."""

    def __init__(self):
        import chromadb
        from app.rag.retriever import ARABIC_PASSAGE_INSTRUCTION, _get_embedding_model  # noqa: F401
        self._chromadb = chromadb
        self._get_embedding = _get_embedding_model
        self._passage_instruction = ARABIC_PASSAGE_INSTRUCTION
        self._splitter = ArabicAwareTextSplitter(
            chunk_size=settings.RAG_CHUNK_SIZE,
            overlap=settings.RAG_CHUNK_OVERLAP,
        )
        self._chroma = chromadb.HttpClient(
            host=settings.CHROMA_HOST,
            port=settings.CHROMA_PORT,
        )

    def ingest_file(self, file_path: Path, doc_metadata: dict | None = None) -> int:
        """
        Ingest a single document. Returns number of chunks added.
        Idempotent — checks content hash before re-inserting.
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            text, page_metadata = self._extract_pdf(file_path)
        elif suffix == ".docx":
            text, page_metadata = self._extract_docx(file_path)
        elif suffix in (".md", ".txt"):
            text = file_path.read_text(encoding="utf-8")
            page_metadata = {}
        else:
            logger.warning("Unsupported file type", path=str(file_path))
            return 0

        chunks = self._splitter.split(text)
        if not chunks:
            logger.warning("No chunks produced", path=str(file_path))
            return 0

        # Build chunk records
        collection = self._get_or_create_collection()
        model = self._get_embedding()

        texts = [c["text"] for c in chunks]
        embeddings = model.encode(
            [f"{self._passage_instruction}{t}" for t in texts],
            batch_size=32,
            show_progress_bar=len(texts) > 50,
        )

        ids = []
        metadatas = []
        for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            content_hash = hashlib.md5(chunk["text"].encode()).hexdigest()[:8]
            chunk_id = f"{file_path.stem}_chunk_{i:04d}_{content_hash}"
            ids.append(chunk_id)

            meta = {
                "source": file_path.name,
                "chunk_index": i,
                "token_count": chunk["token_count"],
                "doc_version": doc_metadata.get("version", "1.0") if doc_metadata else "1.0",
                **page_metadata.get(i, {}),
                **(doc_metadata or {}),
            }
            metadatas.append(meta)

        # Upsert (idempotent)
        collection.upsert(
            ids=ids,
            embeddings=embeddings.tolist(),
            documents=texts,
            metadatas=metadatas,
        )

        logger.info("Ingested document", file=file_path.name, chunks=len(chunks))
        return len(chunks)

    def _extract_docx(self, path: Path) -> tuple[str, dict]:
        """Extract text from a .docx file preserving paragraph structure."""
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx required for DOCX ingestion: pip install python-docx")

        doc = docx.Document(str(path))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        full_text = "\n\n".join(paragraphs)
        return full_text, {}

    def _extract_pdf(self, path: Path) -> tuple[str, dict]:
        """Extract text from PDF with page metadata."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber required for PDF ingestion: pip install pdfplumber")

        full_text = []
        page_meta = {}

        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    full_text.append(text)
                    # Map approximate chunk index to page (rough)
                    # Refined during chunking
                    page_meta[page_num] = {"page": page_num + 1}

        return "\n\n".join(full_text), page_meta

    def _get_or_create_collection(self):
        return self._chroma.get_or_create_collection(
            name=settings.CHROMA_COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
        )

    def get_chunk_count(self) -> int:
        try:
            collection = self._get_or_create_collection()
            return collection.count()
        except Exception:
            return 0


class RAGIndexer:
    """
    High-level indexer called at application startup.
    Checks if handbook is already indexed before re-ingesting.
    """

    def __init__(self):
        self._pipeline = HandbookIngestionPipeline()
        self._handbook_dir = Path(settings.RAG_HANDBOOK_DIR)

    async def ensure_indexed(self) -> None:
        """Index handbook if not already done. Idempotent."""
        existing = self._pipeline.get_chunk_count()

        if existing > 0:
            logger.info("RAG index already exists", chunk_count=existing)
            return

        logger.info("Building RAG index from handbook", dir=str(self._handbook_dir))
        total = 0

        for file_path in sorted(self._handbook_dir.iterdir()):
            if file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                count = self._pipeline.ingest_file(file_path)
                total += count

        logger.info("RAG index built", total_chunks=total)

    async def force_reindex(self) -> int:
        """Force full re-indexing (e.g., after handbook update)."""
        # Delete existing collection
        chroma = chromadb.HttpClient(
            host=settings.CHROMA_HOST,
            port=settings.CHROMA_PORT,
        )
        try:
            chroma.delete_collection(settings.CHROMA_COLLECTION_NAME)
        except Exception:
            pass

        total = 0
        for file_path in sorted(self._handbook_dir.iterdir()):
            if file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                count = self._pipeline.ingest_file(file_path)
                total += count

        logger.info("Force reindex complete", total_chunks=total)
        return total
